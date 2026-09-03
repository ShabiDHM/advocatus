# FILE: backend/app/services/conversion_service.py
# PHOENIX PROTOCOL - CONVERSION SERVICE V13.0 (SAFE TXT/OFFICE PIPELINE & COLLISION-FREE UUID)

import logging
import os
import subprocess
import tempfile
import shutil
import uuid
from PIL import Image

logger = logging.getLogger(__name__)


def _text_to_pdf_fallback(text_content: str, dest_pdf_path: str, title: str = "Dokument") -> str:
    """Konvertues i pastër në Python për tekst në PDF pa pasur nevojë për LibreOffice."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import simpleSplit

        c = canvas.Canvas(dest_pdf_path, pagesize=letter)
        width, height = letter
        margin = 54
        usable_width = width - (2 * margin)

        # Header Title
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, height - margin, title[:60])
        c.line(margin, height - margin - 6, width - margin, height - margin - 6)

        c.setFont("Helvetica", 10)
        y = height - margin - 30

        lines = text_content.split('\n')
        for line in lines:
            if not line.strip():
                y -= 10
                continue
            
            wrapped_lines = simpleSplit(line, "Helvetica", 10, usable_width)
            for wline in wrapped_lines:
                if y < margin:
                    c.showPage()
                    c.setFont("Helvetica", 10)
                    y = height - margin
                c.drawString(margin, y, wline)
                y -= 14

        c.save()
        return dest_pdf_path
    except Exception as e:
        logger.warning(f"ReportLab PDF fallback failed: {e}")
        return ""


def _extract_docx_text_python(source_path: str) -> str:
    """Nxjerr tekstin nga skedarët .docx duke përdorur python-docx."""
    try:
        import docx
        doc = docx.Document(source_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                full_text.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return "\n".join(full_text)
    except Exception as e:
        logger.warning(f"python-docx text extraction failed: {e}")
        return ""


def _extract_txt_file_python(source_path: str) -> str:
    """Lexon me siguri tekstin e plotë të skedarëve .txt me enkoding të pastër."""
    try:
        with open(source_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    except Exception as e:
        logger.warning(f"TXT read failed: {e}")
        return ""


def convert_to_pdf(source_path: str) -> str:
    """
    Konverton çdo format (PDF, DOCX, DOC, TXT, JPG, PNG) në PDF preview.
    Përdor LibreOffice me fallback automatik në ReportLab dhe Pillow.
    """
    if not os.path.exists(source_path):
        raise FileNotFoundError(f"Source file not found at path: {source_path}")

    file_name, source_ext = os.path.splitext(os.path.basename(source_path))
    ext = source_ext.lower()
    
    # PHOENIX FIX: Përdor UUID për të parandaluar mbishkrimin aksidental të skedarëve me emra të njëjtë
    unique_id = uuid.uuid4().hex[:8]
    output_dir = tempfile.gettempdir()
    dest_pdf_path = os.path.join(output_dir, f"{file_name}_{unique_id}_preview.pdf")

    # --- RASTI 1: ËSHTË TASHMË PDF ---
    if ext == '.pdf':
        logger.info(f"Source is already PDF. Copying...")
        shutil.copy2(source_path, dest_pdf_path)
        return dest_pdf_path

    # --- RASTI 2: FOTO NË PDF (JPG, PNG, TIFF, BMP) ---
    if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp']:
        try:
            logger.info(f"Converting Image to PDF: {source_path}")
            image = Image.open(source_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            image.save(dest_pdf_path, "PDF", resolution=150.0)
            logger.info(f"Successfully converted Image to PDF: {dest_pdf_path}")
            return dest_pdf_path
        except Exception as e:
            logger.error(f"Image conversion failed: {e}", exc_info=True)

    # --- RASTI 3: SKEDARËT E ZYRËS (DOCX, DOC) ME LIBREOFFICE ---
    logger.info(f"Initiating Office conversion for '{file_name}{source_ext}'.")
    command = [
        "soffice",
        "--headless",
        "--convert-to", "pdf:writer_pdf_Export",
        "--outdir", output_dir,
        source_path,
    ]

    try:
        process = subprocess.run(
            command,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60
        )
        expected_output_path = os.path.join(output_dir, f"{file_name}.pdf")

        if process.returncode == 0 and os.path.exists(expected_output_path) and os.path.getsize(expected_output_path) > 0:
            shutil.move(expected_output_path, dest_pdf_path)
            logger.info(f"Successfully converted Office Doc to PDF via LibreOffice.")
            return dest_pdf_path
        else:
            logger.warning("LibreOffice not available or failed. Invoking pure-Python fallback...")
    except Exception as libre_err:
        logger.warning(f"LibreOffice execution skipped ({libre_err}). Invoking pure-Python fallback...")

    # --- RASTI 4: PURE PYTHON FALLBACK PËR WORD DHE TXT ---
    if ext == '.txt':
        # PHOENIX FIX: Lexohet teksti real pa u përplasur me python-docx
        extracted_text = _extract_txt_file_python(source_path)
        fallback_pdf = _text_to_pdf_fallback(extracted_text, dest_pdf_path, title=file_name)
        if fallback_pdf and os.path.exists(fallback_pdf):
            logger.info("Successfully converted TXT to PDF using Pure-Python fallback!")
            return dest_pdf_path

    if ext in ['.docx', '.doc']:
        extracted_text = _extract_docx_text_python(source_path)
        if not extracted_text:
            extracted_text = f"Dokument i ngarkuar: {file_name}{source_ext}"
            
        fallback_pdf = _text_to_pdf_fallback(extracted_text, dest_pdf_path, title=file_name)
        if fallback_pdf and os.path.exists(fallback_pdf):
            logger.info("Successfully converted Word doc to PDF using Pure-Python fallback!")
            return dest_pdf_path

    raise RuntimeError(f"Could not convert '{file_name}{source_ext}' to PDF preview.")
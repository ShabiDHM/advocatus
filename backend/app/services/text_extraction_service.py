# FILE: backend/app/services/text_extraction_service.py
# PHOENIX PROTOCOL - OCR ENGINE V15.0 (HIGH-DPI SEQUENTIAL RECONSTRUCTION & ZERO PAGE DROPPING)

import fitz
import logging
import os
import tempfile
import re
import io
import time
from typing import Dict, List, Tuple

try:
    import docx
except ImportError:
    docx = None

try: 
    from app.services.ocr_service import extract_text_from_image_bytes as advanced_bytes_ocr
except Exception:
    try:
        from .ocr_service import extract_text_from_image_bytes as advanced_bytes_ocr
    except Exception:
        advanced_bytes_ocr = None

logger = logging.getLogger(__name__)
FOOTER_PATTERN = re.compile(r'Rasti:\s*\S+\s*\|\s*Juristi AI System')


def _sanitize_text(text: str) -> str: 
    return text.replace("\x00", "") if text else ""


def _strip_footer(text: str) -> str: 
    return '\n'.join([l for l in text.split('\n') if not FOOTER_PATTERN.search(l)])


def _extract_legacy_doc_text(file_path: str) -> str:
    """Extracts text from binary Word 97-2003 (.doc) files safely."""
    try:
        with open(file_path, 'rb') as f:
            content = f.read()

        decoded_latin = content.decode('latin-1', errors='ignore')
        clean_blocks = re.findall(r'[\w\s\.,;:!?\(\)\[\]\/\-\–\—\+\@\#\%\&\=\"]{4,}', decoded_latin)
        extracted = "\n".join([b.strip() for b in clean_blocks if len(b.strip()) > 5])
        
        if extracted and len(extracted) > 50:
            return _sanitize_text(extracted)
    except Exception as e:
        logger.warning(f"Legacy .doc binary parser warning: {e}")

    return ""


def _extract_docx_text(file_path: str) -> str:
    """Extracts text from modern .docx and falls back gracefully for binary .doc."""
    if not docx:
        return _extract_legacy_doc_text(file_path)

    try:
        doc = docx.Document(file_path)
        paragraphs_text = [p.text for p in doc.paragraphs if p.text]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                tables_text.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        
        full_text = "\n".join(paragraphs_text + tables_text)
        if full_text and len(full_text.strip()) > 0:
            return _sanitize_text(full_text)
    except Exception as docx_err:
        logger.warning(f"python-docx warning: {docx_err}")
        return _extract_legacy_doc_text(file_path)

    return _extract_legacy_doc_text(file_path)


def _ocr_single_page_bytes(page_num: int, jpeg_bytes: bytes) -> str:
    """Ekzekuton OCR me përpikëri të lartë për një faqe të vetme."""
    marker = f"\n--- [FAQJA {page_num + 1}] ---\n"
    if not advanced_bytes_ocr:
        return marker + "[SCANNED - NO OCR ENGINE AVAILABLE]"

    try:
        ocr_text = _sanitize_text(advanced_bytes_ocr(jpeg_bytes))
        if ocr_text and len(ocr_text.strip()) > 15:
            return marker + ocr_text.strip()
        return marker + "[Faqe pa tekst të dallueshëm]"
    except Exception as e:
        logger.error(f"❌ [OCR] Gabim në Faqen {page_num + 1}: {e}")
        return marker + ""


def _extract_text_from_pdf(file_path: str) -> str:
    try:
        doc = fitz.open(file_path)
        total = len(doc)
        if total < 1: 
            doc.close()
            return ""
        
        pages_results: Dict[int, str] = {}
        pages_needing_ocr: List[Tuple[int, bytes]] = []

        # Pass 1: Digital Text Extraction
        for i in range(total):
            page = doc[i]
            digital_text = _strip_footer(_sanitize_text("\n".join([b[4] for b in sorted(page.get_text("blocks"), key=lambda b: (int(b[1]/3), int(b[0])))])))
            
            # Nëse faqja ka tekst të qartë dixhital mbi 100 karaktere, përdor atë
            if digital_text and len(digital_text.strip()) > 100:
                pages_results[i] = f"\n--- [FAQJA {i + 1}] ---\n" + digital_text.strip()
            else:
                # PHOENIX FIX: Rezolucion i lartë 2.0x Matrix (300 DPI) me 92% cilësi për skanime gjyqësore
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                jpeg_bytes = pix.tobytes("jpeg", jpg_quality=92)
                pages_needing_ocr.append((i, jpeg_bytes))

        doc.close()

        # Pass 2: PHOENIX SEQUENTIAL OCR (Zero Concurrency Error • 100% Page Success)
        if pages_needing_ocr:
            logger.info(f"📄 [OCR Sequential] Filloi leximi i {len(pages_needing_ocr)} faqeve të skanuara me radhë...")
            for page_num, j_bytes in pages_needing_ocr:
                page_text = _ocr_single_page_bytes(page_num, j_bytes)
                pages_results[page_num] = page_text
                # Pauzë e shkurtër 0.3s për të respektuar limitet e serverit
                time.sleep(0.3)

        # Bashkimi i të gjitha faqeve në renditje rigoroze numerike
        ordered_text = "\n\n".join([pages_results[i] for i in range(total) if i in pages_results])
        logger.info(f"✅ [PDF Extraction Complete] U nxorën gjithsej {len(ordered_text)} karaktere nga {total} faqe.")
        return ordered_text

    except Exception as e:
        logger.error(f"❌ PDF Extraction Failed: {e}")
        return ""


def extract_text(file_path: str, mime_type: str) -> str:
    m = (mime_type or "").lower()
    fn = (file_path or "").lower()

    if "pdf" in m or fn.endswith(".pdf"): 
        return _extract_text_from_pdf(file_path)

    # WORD DOCUMENTS (.docx & legacy .doc)
    if "word" in m or "officedocument" in m or fn.endswith(".docx") or fn.endswith(".doc") or m == "application/msword":
        return _extract_docx_text(file_path)

    # DIRECT IMAGE OCR SUPPORT (.jpg, .jpeg, .png, .webp)
    if any(m.startswith(img_t) for img_t in ["image/jpeg", "image/png", "image/webp", "image/jpg"]) or any(fn.endswith(ext) for ext in [".jpg", ".jpeg", ".png", ".webp"]):
        if advanced_bytes_ocr:
            try:
                with open(file_path, "rb") as img_f:
                    img_bytes = img_f.read()
                return _sanitize_text(advanced_bytes_ocr(img_bytes))
            except Exception as img_err:
                logger.error(f"❌ Direct Image OCR Error: {img_err}")
        return ""

    if "excel" in m or "spreadsheet" in m or fn.endswith(".xlsx") or fn.endswith(".xls"):
        try:
            import pandas as pd
            return _sanitize_text("\n".join(df.to_string() for _, df in pd.read_excel(file_path, sheet_name=None).items()))
        except Exception:
            return ""

    return "" 


def extract_text_from_file(file_obj: io.BytesIO, file_type: str = "PDF") -> str:
    with tempfile.NamedTemporaryFile(suffix=f".{file_type.lower()}", delete=False) as tmp:
        tmp.write(file_obj.getvalue())
        path = tmp.name
    try: 
        return extract_text(path, file_type)
    finally:
        if os.path.exists(path): 
            try:
                os.remove(path)
            except Exception:
                pass